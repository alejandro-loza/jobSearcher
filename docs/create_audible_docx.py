#!/usr/bin/env venv/bin/python3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Crear documento
doc = Document()

# Título principal
title = doc.add_heading("AUDIBLE INTERVIEW CHEATSHEET", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
meta = doc.add_paragraph("Fecha: 10 de abril 2026")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Ubicación: Remote (Zoom)")
doc.add_paragraph("")

# Schedule Table
doc.add_heading("Schedule - 10 de abril 2026", 1)

schedule_data = [
    ("Hora (CDMX)", "Entrevistador", "Tema", "Principios"),
    ("9:30-9:45am", "David Han", "Intro", "-"),
    ("9:45-10:00am", "Break", "-", "-"),
    (
        "10:00-11:00am",
        "Matt Love & Kshitij Shah",
        "Coding: DSA",
        "Articulate The Possible",
    ),
    (
        "11:00-12:00pm",
        "Alfonso Lopez",
        "Coding: Problem Solving",
        "Be Customer Obsessed",
    ),
    ("12:00-1:00pm", "Chaitra Ramdas", "System Design", "Activate Caring + Study/Tech"),
    ("1:00-2:00pm", "Poorva Karunakaran", "Coding: Logical", "Imagine and Invent"),
    ("2:00-2:15pm", "Nathan Shen", "Cierre", "-"),
]

table = doc.add_table(rows=len(schedule_data), cols=4)
for i, row_data in enumerate(schedule_data):
    row_cells = table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]

doc.add_paragraph("")

# Principios de Audible
doc.add_heading("Principios de Audible (Español + Inglés)", 1)

principios = [
    ("Activate Caring", "Activa el Cuidado", "Work human-to-human, no egos"),
    (
        "Articulate The Possible",
        "Articula Lo Posible",
        "Crystallize vision, move with mission",
    ),
    ("Imagine and Invent", "Imagina e Inventa", "Embrace ambiguity, simplify early"),
    (
        "Be Customer Obsessed",
        "Está Obsesionado con el Cliente",
        "Active listening, customer dependency is honor",
    ),
    (
        "Study and Draw Inspiration",
        "Estudia y Obtén Inspiración",
        "Operate at cutting edge",
    ),
]

for eng, esp, desc in principios:
    p = doc.add_paragraph(f"{eng} / {esp}")
    run = p.add_run(desc)
    run.bold = True

doc.add_paragraph("")

# STAR Framework
doc.add_heading("Framework STAR", 1)
doc.add_paragraph("Estructura:")
star = [
    "S = Situation (Situación) - Contexto",
    "T = Task (Tarea) - Tu responsabilidad",
    "A = Action (Acción) - Qué hiciste TÚ específicamente",
    "R = Result (Resultado) - Resultado medible",
]
for item in star:
    p = doc.add_paragraph(f"{item}")
    p.runs[0].bold = True

# STAR Stories - Interview 1
doc.add_heading("Interview 1: Matt Love & Kshitij Shah (10:00-11:00am)", 1)
doc.add_paragraph("Principio: Articulate The Possible + Move Fast")
doc.add_paragraph("")

# Story 1 - Kubo Financiero
doc.add_heading("STAR Story 1: Migración de archivos en Kubo Financiero", 2)
story1_esp = """
SITUACIÓN: En Kubo Financiero (2023), teníamos una aplicación monolítica en servidor local que fallaba frecuentemente por espacio en disco.

TAREA: Necesitaba solucionar el problema de estabilidad del servidor.

ACCIÓN: Organicé un brainstorming con mi equipo. Un compañero propuso migrar archivos a GCP, pero requería construir librería Node.js. Tomé esa idea como inspiración y propuse migrar a GCP pero mapear paths en la nube a los mismos paths locales que el sistema ya esperaba.

RESULTADO: Resolvimos los crashes con cambios mínimos, redujimos uso de recursos, enviamos mucho más rápido.
"""

doc.add_paragraph("Español:")
doc.add_paragraph(story1_esp)

story1_eng = """
SITUATION: At Kubo Financiero (2023), monolithic app on local server crashing due to disk issues.

TASK: Solve server stability issue.

ACTION: Brainstormed with team. Proposed GCP migration with path mapping to existing local structure.

RESULT: Solved crashes with minimal code changes, reduced server usage, shipped faster than original proposal.
"""

doc.add_paragraph("English:")
doc.add_paragraph(story1_eng)

# Story 2 - Clickonero
doc.add_heading("STAR Story 2: Clickonero Hot Sale", 2)
story2_esp = """
SITUACIÓN: En Clickonero, durante Hot Sale, el sistema se cayó.

TAREA: Restaurar el servicio.

ACCIÓN: Me comuniqué con ex-tech lead de otra empresa. Aunque ya no trabajaba con nosotros, respondió a una llamada rápida y nos ayudó a identificar causa raíz.

RESULTADO: Resolvimos incidente en medio día. Sin esa llamada, habría tomado un día completo.
"""

doc.add_paragraph("Español:")
doc.add_paragraph(story2_esp)

story2_eng = """
SITUATION: At Clickonero, during Hot Sale, system went down.

TASK: Restore service.

ACTION: Reached out to former tech lead. He jumped on call, helped identify root cause faster.

RESULT: Resolved incident in half a day.
"""

doc.add_paragraph("English:")
doc.add_paragraph(story2_eng)

# Coding Problems Reference
doc.add_heading("Coding Problems Reference", 1)

# Problem 1 - Two Sum
doc.add_heading("1. Audiobook Trip Recommender (Two Sum)", 2)
doc.add_paragraph(
    "Pregunta: Dada duración de viaje y lista de libros, encontrar par que sume exactamente."
)

code1 = """public static List<String> recommendBooks(double tripDuration, Map<String, Double> books) {
    Map<Double, String> seen = new HashMap<>();

    for (Map.Entry<String, Double> entry : books.entrySet()) {
        double complement = tripDuration - entry.getValue();
        if (seen.containsKey(complement)) {
            return Arrays.asList(seen.get(complement), entry.getKey());
        }
        seen.put(entry.getValue(), entry.getKey());
    }
    return Collections.emptyList();
}"""

doc.add_paragraph("Solución óptima (O(n)):")
doc.add_paragraph(code1)

# Problem 2 - Credit Management
doc.add_heading("2. Credit Management System (FIFO + Expiration)", 2)
doc.add_paragraph(
    "Pregunta: Sistema de créditos con expiración de 1 año, consumo FIFO."
)

code2 = """class CreditManager {
    private Deque<double[]> credits = new ArrayDeque<>();

    public void addCredit(LocalDate date, double amount) {
        LocalDate expiry = date.plusYears(1);
        credits.addLast(new double[]{amount, expiry.toEpochDay()});
    }

    public double getBalance(LocalDate queryDate) {
        double balance = 0;
        long queryEpoch = queryDate.toEpochDay();

        for (double[] credit : credits) {
            if (credit[1] >= queryEpoch) {
                balance += credit[0];
            }
        }
        return balance;
    }
}"""

doc.add_paragraph("Solución con Deque (O(n)):")
doc.add_paragraph(code2)

# Amazon Locker System Design
doc.add_heading("Amazon Locker System Design", 1)

doc.add_paragraph("Arquitectura:")
doc.add_paragraph("Client → API Gateway → Locker Service → DB (PostgreSQL)")
doc.add_paragraph("Fan-out: SQS/Kafka → Feed Cache (Redis)")

doc.add_paragraph("")
doc.add_paragraph("Clarifying Questions:")
questions = [
    "¿Cuántos usuarios diarios? 10k vs 10M",
    "¿Real-time o eventual consistency?",
    "¿Con media o solo texto?",
    "¿Con búsqueda o solo feed?",
]
for q in questions:
    p = doc.add_paragraph(f"• {q}")

doc.add_paragraph("")
doc.add_paragraph("Trade-offs:")
tradeoffs = [
    (
        "Feed",
        "Push (fast reads, expensive for celebrities) vs Pull (scalable for celebrities)",
    ),
    ("DB", "SQL (ACID, joins) vs NoSQL (scale, eventual consistency)"),
]
for decision, tradeoff in tradeoffs:
    p = doc.add_paragraph(f"{decision}: {tradeoff}")

# Tips Generales
doc.add_heading("Tips Generales", 1)

tips = [
    "Máximo 2 minutos por historia STAR",
    "En coding: Talk out loud - nunca silencio",
    "Empieza con brute force, luego optimiza",
    "Menciona trade-offs proactivamente",
    "Code readability > Cleverness",
]
for tip in tips:
    p = doc.add_paragraph(f"✓ {tip}")

# Checklist
doc.add_heading("Checklist Final", 1)
checklist = [
    "18 STAR stories memorizadas",
    "3 soluciones de coding memorizadas",
    "Arquitectura de Amazon Locker memorizada",
    "CV actualizado con proyectos de AI/LLM",
    "LinkedIn actualizado con experiencia reciente",
    "Práctica de mock interviews (2 sesiones)",
]
for i, item in enumerate(checklist, 1):
    doc.add_paragraph(f"[ ] {i}. {item}")

# Footer
footer = doc.add_paragraph("© 2026 - Alejandro Hernandez Loza")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("¡Buena suerte en las entrevistas de Audible!")
run.bold = True

# Guardar
output_path = "docs/AUDIBLE_INTERVIEW_FINAL.docx"
doc.save(output_path)
print(f"Documento DOCX creado: {output_path}")
print(f"Tamaño del archivo: {os.path.getsize(output_path)} bytes")
