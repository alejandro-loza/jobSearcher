"""
Genera el cheatsheet completo de Audible interview en formato .docx
Basado en: principios Audible, historias STAR de Alejandro, problemas probables
investigados desde Glassdoor, LeetCode Amazon top questions, audiblecareers.com
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
    p.runs[0].font.size = Pt(size)
    return p


def add_code(doc, code_text, size=8):
    """Add code block with very light background for readability.

    Keeps all lines of the code together on a single page when possible.
    """
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # Tight line spacing so code fits better on one page
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    # Keep all lines of this paragraph on the same page
    p.paragraph_format.keep_together = True
    # Very subtle off-white background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FAFAFA')
    pPr.append(shd)
    # Add a light border via paragraph properties
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        pBdr.append(b)
    pPr.append(pBdr)
    return p


def read_java_file(path):
    """Read a Java file and return clean code without package/class wrapper comments."""
    try:
        with open(path, 'r') as f:
            return f.read()
    except Exception:
        return ''


def add_code_example(doc, title, file_path, description='', color_hex='1F4E79'):
    """Add a code example section with title, description and code block.

    Title and description are set with keep_with_next so they stay on the same
    page as the code block that follows.
    """
    p = doc.add_paragraph()
    run = p.add_run('→ ' + title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    if description:
        desc_p = add_para(doc, description, size=9)
        desc_p.paragraph_format.keep_with_next = True
        desc_p.paragraph_format.space_after = Pt(2)
    code = read_java_file(file_path)
    if code:
        add_code(doc, code, size=8)


# ============================================================
# STAR story helpers — Spanish + English versions side by side
# ============================================================

def add_star_title(doc, text, color_hex='1F4E79'):
    """Story heading (principle tag + short title)."""
    rgb = RGBColor.from_string(color_hex)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = rgb
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_star_why(doc, text, color_hex='1F4E79'):
    """Italic line explaining why the story matches the principle."""
    rgb = RGBColor.from_string(color_hex)
    p = doc.add_paragraph()
    run = p.add_run('Why this fits: ')
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = rgb
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(2)
    return p


def add_star_lang_label(doc, text, color_hex='1F4E79'):
    """Small label like 'ES' or 'EN' before a story block."""
    rgb = RGBColor.from_string(color_hex)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = rgb
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    return p


def add_star_line(doc, letter, text):
    """One STAR line (S/T/A/R). Letter bold, text regular."""
    p = doc.add_paragraph()
    run = p.add_run(letter + ': ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    return p


def add_star_story(doc, title, why_it_fits, star_es, star_en, color_hex='1F4E79'):
    """Add a full STAR story block with Spanish + English versions.

    star_es / star_en are dicts with keys: S, T, A, R (some may be missing).
    """
    add_star_title(doc, title, color_hex=color_hex)
    add_star_why(doc, why_it_fits, color_hex=color_hex)

    add_star_lang_label(doc, 'ES', color_hex=color_hex)
    for letter in ('S', 'T', 'A', 'R'):
        if letter in star_es and star_es[letter]:
            add_star_line(doc, letter, star_es[letter])

    add_star_lang_label(doc, 'EN', color_hex=color_hex)
    for letter in ('S', 'T', 'A', 'R'):
        if letter in star_en and star_en[letter]:
            add_star_line(doc, letter, star_en[letter])


def add_section_title(doc, text, color_hex='1F4E79'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()

    # Margenes pequeños para maximizar contenido
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Font default
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ============================================================
    # PORTADA
    # ============================================================
    title = doc.add_heading('AUDIBLE INTERVIEW CHEATSHEET', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Viernes 10 de Abril 2026  |  Alejandro Hernandez Loza')
    run.bold = True
    run.font.size = Pt(12)

    # Info Zoom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Zoom: ').bold = True
    p.add_run('https://amazon.zoom.us/j/92735802948\n')
    p.add_run('Meeting ID: ').bold = True
    p.add_run('9273 580 2948   ')
    p.add_run('Password: ').bold = True
    p.add_run('12618815')

    doc.add_paragraph()

    # Schedule table
    add_section_title(doc, 'SCHEDULE DEL DIA')
    sched = doc.add_table(rows=1, cols=3)
    sched.style = 'Light Grid Accent 1'
    hdr = sched.rows[0].cells
    hdr[0].text = 'Hora'
    hdr[1].text = 'Entrevistador'
    hdr[2].text = 'Enfoque'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    schedule_rows = [
        ('9:30-9:45', 'David Han (Recruiting)', 'Intro/bienvenida'),
        ('10:00-11:00', 'Matt Love + Kshitij Shah (SDE II)', 'Coding DSA + Articulate The Possible'),
        ('11:00-12:00', 'Alfonso Lopez (SDM)', 'Problem Solving + Be Customer Obsessed'),
        ('12:00-13:00', 'Chaitra Ramdas (SDE II)', 'System Design + Activate Caring + Study & Draw'),
        ('13:00-14:00', 'Poorva Karunakaran (SDE II)', 'Logical & Maintainable + Imagine & Invent'),
        ('14:00-14:15', 'Nathan Shen (Recruiter)', 'Cierre'),
    ]
    for row in schedule_rows:
        cells = sched.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph()

    # Audible Principles summary
    add_section_title(doc, 'AUDIBLE PEOPLE PRINCIPLES')
    principles = [
        ('Articulate The Possible & Move Fast', 'Cristaliza visiones con palabras bien escogidas, usa data para hipótesis, muévete rápido con propósito misionero.'),
        ('Be Customer Obsessed', 'Escucha voces del cliente incluso cuando no hablan. Su dependencia es un honor; supera expectativas.'),
        ('Activate Caring', 'Hombro a hombro, sin egos, altos estándares con amabilidad. La gente olvida palabras, no cómo los hiciste sentir.'),
        ('Study & Draw Inspiration', 'Opera en la vanguardia de artes y ciencias. Aprende de cultura, tecnología emergente, ciencias humanas.'),
        ('Imagine & Invent Before They Ask', 'Abraza ambigüedad, simplifica complejidad. Haz que el cliente sienta el poder antes de tener que pensarlo.'),
    ]
    for name, desc in principles:
        p = doc.add_paragraph()
        p.add_run('• ' + name + ': ').bold = True
        p.add_run(desc)

    add_page_break(doc)

    # ============================================================
    # INTERVIEW 1 - MATT & KSHITIJ
    # ============================================================
    add_heading(doc, 'INTERVIEW 1: Matt Love + Kshitij Shah', level=1, color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, '10:00-11:00am  |  Coding (DSA) + Articulate The Possible & Move Fast', bold=True, size=11)
    doc.add_paragraph()

    # Principio
    add_section_title(doc, 'PRINCIPIO: Articulate The Possible & Move Fast')
    add_para(doc, '"Crystallize a vision with a few well-chosen words. Use data to support your hypotheses and move with missionary purpose."', size=9)
    add_para(doc, 'Lo que evalúan:', bold=True, size=10)
    eval_1 = [
        'Capacidad de INSPIRAR con palabras — no jerga, sino claridad que mueve a la acción',
        'Uso de datos e hipótesis para respaldar decisiones, incluso en incertidumbre',
        'Velocidad de ejecución — no esperas permiso, actúas con propósito',
        'Colaboración: obtienes mejores ideas del equipo y las integras, no impones',
        'Agilidad: cambias de dirección rápido cuando los hechos cambian',
    ]
    for e in eval_1:
        add_bullet(doc, e, size=9)
    add_para(doc, 'Frases que disparan señales positivas en el entrevistador:', bold=True, size=10)
    signals_1 = [
        '"I aligned the team around a clear goal: X, by when, and how success looks"',
        '"I used data to validate the hypothesis before committing resources"',
        '"I activated my network — reached out to [person] who helped us move 10x faster"',
        '"I didn\'t wait for direction — I saw the window of opportunity and moved"',
        '"I solicited three team ideas, synthesized them, and proposed a combined approach"',
    ]
    for s in signals_1:
        add_bullet(doc, s, size=9)
    add_para(doc, 'Red flags a EVITAR:', bold=True, size=10)
    red_1 = [
        'Historias donde solo tú hiciste todo — sin equipo, sin colaboración',
        'Sin datos ni métricas concretas para validar tu hipótesis',
        'Proceso muy lento, muchas reuniones, sin sentido de urgencia',
        'Hablar en "nosotros" sin clarificar TU contribución específica',
    ]
    for r in red_1:
        add_bullet(doc, r, size=9)

    # Behavioral questions
    add_section_title(doc, 'PREGUNTAS BEHAVIORAL PROBABLES')
    questions_1 = [
        '"Give me an example of when you solicited ideas from your team and used those to come to a better solution than one you would have developed yourself."',
        '"Tell me about a time you rolled up your sleeves to get something done, even if it wasnt technically your piece of work."',
        '"Tell me about a time you needed to quickly get your team up to speed on new information."',
        '"Give me an example of a time you leveraged your internal network to deliver a better or faster result."',
        '"Tell me about a time you had to inspire others with a vision."',
    ]
    for q in questions_1:
        add_bullet(doc, q)

    # STAR Stories
    add_section_title(doc, 'STAR STORIES PREPARADAS')

    add_para(doc, 'STORY 1 — Kubo Financiero: Migración de archivos a GCP', bold=True, size=11, color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, 'S: En Kubo Financiero, el monolito on-prem crasheaba frecuentemente por falta de espacio en disco.')
    add_para(doc, 'T: Mi tarea era resolver el problema con el equipo.')
    add_para(doc, 'A: Un teammate propuso migrar a GCP + library Node.js para integrar con el monolito JSP — requería semanas. Yo tomé la idea como inspiración y propuse: migrar archivos a GCP pero mapear los paths de la nube a los mismos paths locales. Los archivos se descargan on-demand de forma transparente — cero cambios al frontend.')
    add_para(doc, 'R: Resolvimos los crashes con cambios mínimos, redujimos uso de recursos del servidor y entregamos mucho más rápido que la propuesta original.')

    doc.add_paragraph()
    add_para(doc, 'STORY 2 — Clickonero Hot Sale (rolled up sleeves)', bold=True, size=11, color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, 'S: En Clickonero durante Hot Sale — uno de los días más ocupados del año — el sistema cayó. El arquitecto estaba viajando y era inaccesible.')
    add_para(doc, 'T: Mi equipo tenía que encontrar y arreglar el problema, pero no era técnicamente mi pieza.')
    add_para(doc, 'A: Contacté a nuestro ex tech lead que ya había dejado la empresa. Aunque no trabajaba con nosotros, se conectó a una llamada rápida y nos ayudó a identificar la causa raíz mucho más rápido.')
    add_para(doc, 'R: Resolvimos el incidente en medio día. Sin esa llamada hubiera tomado el día completo — durante Hot Sale eso hubiera significado pérdidas significativas.')

    # Coding problems
    add_section_title(doc, 'PROBLEMAS DE CODING PROBABLES (DSA)')
    add_para(doc, 'Basado en: Glassdoor Audible/Amazon, top Amazon LeetCode, audiblecareers.com. Esperar problemas medium-hard.', size=9)

    coding_1 = [
        ('Two Sum / Variants (Audiobook Trip)', 'HashMap complement approach. O(n) time, O(n) space. Brute force O(n²) como punto de partida. Para floats usar epsilon.'),
        ('First Unique Character in Stream', 'Two passes con array[26] de frecuencias. O(n) time, O(1) space. Alternativa: LinkedHashMap preserva orden.'),
        ('Longest Substring Without Repeating Characters', 'Sliding window + HashSet. O(n) time. Dos pointers que se mueven solo hacia adelante.'),
        ('Merge K Sorted Lists/Streams', 'Min-Heap de k elementos. O(n log k) time. Útil para audiobook playlists ordenados.'),
        ('LRU Cache', 'HashMap + Doubly Linked List. O(1) get/put. Caso: cache de audiobooks favoritos.'),
        ('Number of Islands (BFS/DFS)', 'DFS marca visitados. O(m*n) time. Variante: connected components en grafo social.'),
        ('Valid Parentheses', 'Stack. Push opening, pop y compara en closing. O(n).'),
        ('Binary Tree Level Order / Zigzag', 'BFS con queue. Alternar dirección con flag boolean.'),
        ('Binary Tree Maximum Path Sum', 'DFS post-order retornando max gain. Track global max.'),
        ('Find First Number Odd Occurrences', 'XOR de todo el array. a XOR a = 0. Ideal para memoria O(1).'),
        ('Reverse String / Linked List In Place', 'Two pointers o recursión. Classic warmup.'),
        ('Top K Frequent Elements', 'HashMap + Heap O(n log k) o Bucket Sort O(n).'),
        ('Subarray Sum Equals K', 'Prefix sum + HashMap. O(n) time.'),
        ('Word Ladder', 'BFS en grafo de palabras. Útil para audiobook recommendation chains.'),
        ('Course Schedule', 'Topological sort (Kahn o DFS). Detección de ciclos.'),
    ]
    for name, desc in coding_1:
        p = doc.add_paragraph()
        p.add_run('• ' + name + ': ').bold = True
        r = p.add_run(desc)
        r.font.size = Pt(9)

    # Protocolo de respuesta
    add_section_title(doc, 'PROTOCOLO DE RESPUESTA EN CODING')
    protocol = [
        '1. CLARIFY: "Let me make sure I understand the problem..." → Repite con tus palabras + 2-3 ejemplos',
        '2. EDGE CASES: empty input, 1 element, duplicates, negatives, overflow',
        '3. BRUTE FORCE: "Let me think of the simplest approach first..." → menciona O(n²) aunque no codifiques',
        '4. OPTIMIZE: "Can I trade space for time with a HashMap?" → articula el trade-off',
        '5. CODE: escribe sintaxis correcta (no pseudocódigo), nombra variables descriptivas',
        '6. TEST: recorre el código con el ejemplo dado, verifica edge cases en voz alta',
        '7. COMPLEXITY: declara Time y Space al final: "This is O(n) time, O(n) space because..."',
    ]
    for line in protocol:
        add_bullet(doc, line, size=9)

    # EJEMPLOS DE CÓDIGO (soluciones ya escritas por Alejandro)
    add_section_title(doc, 'EJEMPLOS DE CODIGO LISTOS (tus soluciones)')
    add_para(doc, 'Estos son patrones que YA tienes resueltos. Úsalos como referencia mental.', size=9)

    add_code_example(
        doc,
        'Two Sum (HashMap complement) — matches el audiobook trip problem',
        '/home/pinky/proyects/leetcode/src/TwoSumSolution.java',
        'Patrón central de Interview 1. O(n) time, O(n) space. Mismo approach aplica al problema de audiobooks + trip duration.',
        color_hex='1F4E79',
    )

    add_code_example(
        doc,
        'Longest Substring Without Repeating (Sliding Window)',
        '/home/pinky/proyects/leetcode/src/LongestSubstringSolution.java',
        'Patrón sliding window con HashMap. Clásico de Amazon/Audible.',
        color_hex='1F4E79',
    )

    add_code_example(
        doc,
        'First Duplicate Character (HashSet)',
        '/home/pinky/proyects/leetcode/src/FirstDuplicateCharSolution.java',
        'Detección de duplicados en un stream. Útil para variantes del problema stream.',
        color_hex='1F4E79',
    )

    add_code_example(
        doc,
        'Maximum Subarray (Kadane algorithm)',
        '/home/pinky/proyects/leetcode/src/MaximumSubarraySolution.java',
        'DP en O(1) space. Patrón de "running sum + max tracker".',
        color_hex='1F4E79',
    )

    add_code_example(
        doc,
        'Container With Most Water (Two Pointers)',
        '/home/pinky/proyects/leetcode/src/ContainerWithMostWaterSolution.java',
        'Two pointers convergentes. Alternativa al HashMap approach cuando el array está ordenado.',
        color_hex='1F4E79',
    )

    # Preguntas para ellos
    add_section_title(doc, 'PREGUNTAS PARA MATT & KSHITIJ (al final)')
    questions_for_them_1 = [
        '"What does a typical day look like for an SDE II on your team?"',
        '"What are the biggest technical challenges the team is solving right now?"',
        '"What does success look like in the first 90 days in this role?"',
        '"How does the team balance shipping fast with code quality?"',
    ]
    for q in questions_for_them_1:
        add_bullet(doc, q)

    add_page_break(doc)

    # ============================================================
    # INTERVIEW 2 - ALFONSO LOPEZ
    # ============================================================
    add_heading(doc, 'INTERVIEW 2: Alfonso Lopez Lopez (SDM)', level=1, color=RGBColor(0x2E, 0x75, 0x4B))
    add_para(doc, '11:00-12:00pm  |  Coding (Problem Solving) + Be Customer Obsessed', bold=True, size=11)
    doc.add_paragraph()

    add_section_title(doc, 'PRINCIPIO: Be Customer Obsessed', color_hex='2E754B')
    add_para(doc, '"That our customers depend on us is an honor. We listen to the voices of our customers even when they don\'t speak."', size=9)
    add_para(doc, 'Lo que evalúan:', bold=True, size=10)
    eval_2 = [
        'Escucha activa REAL — no solo encuestas, sino entender el dolor profundo del usuario',
        'Cambios de dirección guiados por insights de clientes, no por ego o asunciones internas',
        'Disposición a ir "above and beyond" cuando el cliente lo requiere, sin que te lo pidan',
        'Medir el éxito en términos del cliente (satisfacción, adoption, retención), no solo métricas técnicas',
        'Cómo convenciste a stakeholders internos de priorizar al cliente sobre otras presiones',
    ]
    for e in eval_2:
        add_bullet(doc, e, size=9)
    add_para(doc, 'Frases que disparan señales positivas:', bold=True, size=10)
    signals_2 = [
        '"I sat down with the team and listened to their actual workflow, not just their reported pain points"',
        '"The insight changed my initial assumptions entirely — I had to pivot the approach"',
        '"We measured success not in uptime, but in whether the customer\'s problem was truly solved"',
        '"I championed this change internally even when it wasn\'t on the roadmap"',
        '"The customer\'s dependence on us was a responsibility I took seriously, not a burden"',
    ]
    for s in signals_2:
        add_bullet(doc, s, size=9)
    add_para(doc, 'Red flags a EVITAR:', bold=True, size=10)
    red_2 = [
        'Hablar de características técnicas sin conectar con el impacto al usuario final',
        '"El cliente nos pidió X" sin mostrar que tú cuestionaste si eso era lo que realmente necesitaban',
        'Historias donde el resultado es técnico ("bajamos la latencia") sin impacto al cliente',
        'No mencionar cómo validaste que el cliente quedó realmente satisfecho',
    ]
    for r in red_2:
        add_bullet(doc, r, size=9)

    add_section_title(doc, 'PREGUNTAS BEHAVIORAL PROBABLES', color_hex='2E754B')
    questions_2 = [
        '"Give me an example of when you had to change your course of action given a new insight from a customer."',
        '"Tell me about a time you went above and beyond to ensure a customer was delighted."',
        '"Tell me about a time you dove deep into a customer experience that then influenced a business decision."',
        '"Tell me about a time you championed for the customer to make a change within the business."',
        '"Tell me about a time you had to deliver bad news to a customer."',
    ]
    for q in questions_2:
        add_bullet(doc, q)

    add_section_title(doc, 'STAR STORIES PREPARADAS', color_hex='2E754B')

    add_star_story(
        doc,
        'STORY 1 — CMS de Encuestas (Presidencia) [Changed Course from Customer Insight]',
        'Cambiaste el scope basado en escuchar el workflow real del cliente, no en lo inicialmente pedido.',
        {
            'S': 'En Presidencia de México, el equipo de atención ciudadana publicaba encuestas manualmente — pedían al equipo de desarrollo escribir HTML cada vez. Proceso lento y frustrante.',
            'T': 'Inicialmente íbamos a construir solo un formulario simple de envío.',
            'A': 'Me senté con el equipo y escuché su workflow real — no lo que pidieron, sino cómo realmente trabajaban. Entendí que necesitaban autonomía total. Cambié el scope: construimos un CMS self-service donde publicaban encuestas sin tocar código.',
            'R': 'El equipo pasó de esperar días por cada encuesta a publicarlas en minutos. Liberamos tiempo de developers para proyectos de mayor impacto.',
        },
        {
            'S': 'At Mexico\'s Presidential Office, the citizen attention team published polls manually — requesting HTML from developers each time. A frustrating, dependency-heavy process.',
            'T': 'We were initially scoped to build a simple form submission tool.',
            'A': 'I sat down with the team and observed their actual workflow — not just what they had requested, but how they really worked. I realized they needed full autonomy. I pivoted the scope and built a self-service CMS where they could create and publish polls with no developer involvement.',
            'R': 'The team went from waiting days per poll to publishing in minutes on their own. We freed developer capacity for higher-impact projects.',
        },
        color_hex='2E754B',
    )

    add_star_story(
        doc,
        'STORY 2 — Terremoto CDMX (Comodín: Above & Beyond) [Went Above to Delight/Help Customer]',
        'Actuaste con sentido de urgencia extremo en favor del ciudadano, sin seguir procesos normales.',
        {
            'S': 'Trabajaba en Presidencia cuando el terremoto golpeó CDMX. Caos total, rescatistas sin información organizada.',
            'T': 'Un grupo pequeño decidió actuar — publicar un formulario de emergencia en gob.mx donde ciudadanos reportaran ubicación y daños con GPS.',
            'A': 'Nos saltamos security reviews, QA, load testing. El único objetivo: llegar a producción ese día para que los rescatistas tuvieran un mapa en tiempo real de donde la gente necesitaba ayuda.',
            'R': 'El sistema salió el mismo día. Equipos de rescate lo usaron para priorizar zonas. No medimos el éxito en uptime ese día — lo medimos en vidas.',
        },
        {
            'S': 'I was working at Mexico\'s Presidential Office when a major earthquake hit Mexico City. Rescue teams had no organized data on where people needed help.',
            'T': 'A small group of us decided to act — build and publish an emergency form on gob.mx where citizens could report their location, describe damage, and share GPS coordinates.',
            'A': 'We bypassed security reviews, QA, and load testing. The only goal was to get it live the same day so rescue teams had a real-time map of affected areas.',
            'R': 'The system went live the same day. Rescue teams used it to prioritize locations. We didn\'t measure success in uptime that day — we measured it in lives.',
        },
        color_hex='2E754B',
    )

    # Problem solving exercises
    add_section_title(doc, 'EJERCICIOS DE PROBLEM SOLVING PROBABLES', color_hex='2E754B')
    add_para(doc, 'Alfonso es SDM → esperará problemas con componente de diseño/business logic más que puro algoritmo.', size=9)

    ps_problems = [
        ('Audible Credit Management System', 'FIFO Queue + expiración 1 año. Deque<{amount, expiry}>. Balance query itera y suma no-expirados. Clarificar si expira EN la fecha o DESPUES.'),
        ('Rate Limiter', 'Token Bucket o Leaky Bucket. Usar HashMap<userId, Bucket>. O(1) per request. Thread-safe con atomic ops.'),
        ('Design a Parking Lot', 'ParkingLot → List<Floor> → 2D grid de Spots. Strategy pattern para asignación. S/M/L/XL spots. Ticket con timestamp.'),
        ('Design an Elevator System', 'State pattern (Moving Up, Moving Down, Idle). Priority queue de requests. Optimización: pickup más cercano en dirección.'),
        ('Design a Vending Machine', 'State pattern: Idle, HasMoney, Dispensing. Inventory management. Change calculation (greedy coins).'),
        ('Design a Library Management System', 'Book, Member, Librarian, Loan, Fine. Observer pattern para notificaciones de disponibilidad.'),
        ('Design Tic-Tac-Toe', 'Optimización: track row/col/diagonal sums en O(1) check. No re-scan board cada turno.'),
        ('Design a Hit Counter', 'Sliding window con Queue<timestamp> o circular array.'),
        ('Design a Logger Rate Limiter', 'HashMap<message, lastTimestamp>. Rechazar si gap < threshold.'),
        ('File System Design', 'Trie de paths. cd, ls, mkdir, addContent. Composite pattern para files y directories.'),
        ('Design a Stack with getMin() in O(1)', 'Dos stacks: main + minStack. Al push, también push min actual a minStack.'),
        ('LRU Cache (OOP version)', 'HashMap + DoublyLinkedList. Contraste con LFU Cache (frecuencias).'),
        ('Course Schedule / Dependency Resolution', 'Topological sort. Detección de ciclos con 3-color DFS o Kahn.'),
        ('Design a Booking System', 'Calendar con time slots. Conflict detection. Concurrency: optimistic locking o pessimistic.'),
    ]
    for name, desc in ps_problems:
        p = doc.add_paragraph()
        p.add_run('• ' + name + ': ').bold = True
        r = p.add_run(desc)
        r.font.size = Pt(9)

    add_section_title(doc, 'PROTOCOLO PARA PROBLEM SOLVING / OOP', color_hex='2E754B')
    protocol_2 = [
        '1. REQUIREMENTS: functional + non-functional. Aclara scope ("MVP or full system?")',
        '2. USE CASES: identifica actores y acciones principales',
        '3. CLASS DIAGRAM: entidades principales, relaciones, multiplicidad',
        '4. INTERFACES: define contratos antes de implementación',
        '5. DESIGN PATTERNS: menciona cuáles usas y por qué (Strategy, State, Observer, Factory, Composite)',
        '6. SOLID: Single Responsibility, Open/Closed, etc. — nómbralos explícitamente',
        '7. EDGE CASES y concurrency: ¿thread-safe? ¿qué pasa con race conditions?',
    ]
    for line in protocol_2:
        add_bullet(doc, line, size=9)

    # EJEMPLOS DE CÓDIGO para Interview 2
    add_section_title(doc, 'EJEMPLOS DE CODIGO LISTOS (tus soluciones)', color_hex='2E754B')
    add_para(doc, 'Patrones aplicables a problem solving / OOP.', size=9)

    add_code_example(
        doc,
        'Valid Parentheses (Stack pattern)',
        '/home/pinky/proyects/leetcode/src/ValidParenthesesSolution.java',
        'Stack para balance. Mismo patrón aplica a calculator, expression evaluation, XML/HTML parser.',
        color_hex='2E754B',
    )

    # Credit Management System example
    credit_code = '''import java.time.LocalDate;
import java.util.ArrayDeque;
import java.util.Deque;

/**
 * Audible Credit Management System
 * Credits expire 1 year after issuance. FIFO consumption (earliest expiry first).
 */
public class CreditManager {

    private static class Credit {
        double amount;
        LocalDate expiryDate;

        Credit(double amount, LocalDate expiryDate) {
            this.amount = amount;
            this.expiryDate = expiryDate;
        }
    }

    private final Deque<Credit> credits = new ArrayDeque<>();

    public void addCredit(LocalDate issueDate, double amount) {
        if (amount <= 0) {
            throw new IllegalArgumentException("Amount must be positive");
        }
        credits.addLast(new Credit(amount, issueDate.plusYears(1)));
    }

    public double getBalance(LocalDate queryDate) {
        double balance = 0;
        for (Credit c : credits) {
            if (c.expiryDate.isAfter(queryDate)) {
                balance += c.amount;
            }
        }
        return balance;
    }

    public void consume(double amount) {
        if (amount <= 0) return;
        double remaining = amount;
        while (remaining > 0 && !credits.isEmpty()) {
            Credit oldest = credits.peekFirst();
            if (oldest.amount <= remaining) {
                remaining -= oldest.amount;
                credits.pollFirst();
            } else {
                oldest.amount -= remaining;
                remaining = 0;
            }
        }
        if (remaining > 0) {
            throw new IllegalStateException("Insufficient credits");
        }
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run('→ Audible Credit Management (FIFO + Expiration)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0x4B)
    add_para(doc, 'Ejemplo completo del problema Credit System. Deque mantiene orden FIFO. Clarificar con el entrevistador: ¿expira EN la fecha o DESPUES?', size=9)
    add_code(doc, credit_code, size=8)

    # Rate Limiter example
    rate_limiter_code = '''import java.util.concurrent.ConcurrentHashMap;
import java.util.concurrent.atomic.AtomicLong;

/**
 * Token Bucket Rate Limiter — per user
 * Thread-safe with AtomicLong for tokens.
 */
public class TokenBucketRateLimiter {

    private static class Bucket {
        AtomicLong tokens;
        long lastRefillNanos;
        final long capacity;
        final long refillPerSecond;

        Bucket(long capacity, long refillPerSecond) {
            this.capacity = capacity;
            this.refillPerSecond = refillPerSecond;
            this.tokens = new AtomicLong(capacity);
            this.lastRefillNanos = System.nanoTime();
        }
    }

    private final ConcurrentHashMap<String, Bucket> buckets = new ConcurrentHashMap<>();
    private final long capacity;
    private final long refillPerSecond;

    public TokenBucketRateLimiter(long capacity, long refillPerSecond) {
        this.capacity = capacity;
        this.refillPerSecond = refillPerSecond;
    }

    public boolean allow(String userId) {
        Bucket b = buckets.computeIfAbsent(userId,
            k -> new Bucket(capacity, refillPerSecond));
        synchronized (b) {
            long now = System.nanoTime();
            long elapsed = now - b.lastRefillNanos;
            long tokensToAdd = (elapsed * b.refillPerSecond) / 1_000_000_000L;
            if (tokensToAdd > 0) {
                b.tokens.set(Math.min(b.capacity,
                    b.tokens.get() + tokensToAdd));
                b.lastRefillNanos = now;
            }
            if (b.tokens.get() > 0) {
                b.tokens.decrementAndGet();
                return true;
            }
            return false;
        }
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run('→ Rate Limiter (Token Bucket)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0x4B)
    add_para(doc, 'Thread-safe token bucket per-user. Patrón clave para APIs a escala.', size=9)
    add_code(doc, rate_limiter_code, size=8)

    add_section_title(doc, 'PREGUNTAS PARA ALFONSO (SDM)', color_hex='2E754B')
    questions_for_them_2 = [
        '"How do you measure success for your team?"',
        '"What is the engineering culture like? How do you balance autonomy and alignment?"',
        '"How do you handle disagreements on technical decisions within the team?"',
        '"What are the biggest customer problems your team is solving right now?"',
        '"How do you support career growth for engineers on your team?"',
    ]
    for q in questions_for_them_2:
        add_bullet(doc, q)

    add_page_break(doc)

    # ============================================================
    # INTERVIEW 3 - CHAITRA RAMDAS (SYSTEM DESIGN)
    # ============================================================
    add_heading(doc, 'INTERVIEW 3: Chaitra Ramdas (SDE II)', level=1, color=RGBColor(0xBF, 0x8F, 0x00))
    add_para(doc, '12:00-1:00pm  |  System Design + Activate Caring + Study & Draw Inspiration', bold=True, size=11)
    doc.add_paragraph()

    add_section_title(doc, 'PRINCIPIOS: Activate Caring + Study & Draw Inspiration', color_hex='BF8F00')
    add_para(doc, '"People will invariably forget things you say, but they will rarely forget how you make them feel." / "Leaders who stand out operate at the cutting edge of arts and sciences."', size=9)
    add_para(doc, 'ACTIVATE CARING — Lo que evalúan:', bold=True, size=10)
    eval_3a = [
        'Empatía genuina: ayudar a otros sin beneficio propio, sin ego',
        'Altos estándares con amabilidad — no uno o el otro, ambos a la vez',
        'Colaboración real: momentos donde priorizaste el éxito del equipo sobre el tuyo',
        'Relaciones: cómo construiste confianza con colegas más allá del trabajo transaccional',
        'Apoyo a ideas de otros, incluso cuando eran impopulares o no eran tuyas',
    ]
    for e in eval_3a:
        add_bullet(doc, e, size=9)
    add_para(doc, 'STUDY & DRAW — Lo que evalúan:', bold=True, size=10)
    eval_3b = [
        'Curiosidad fuera de tu campo: libros, ciencias, artes, humanidades que influenciaron tu trabajo',
        'Aplicar insights no-obvios a problemas técnicos (ej: diseño industrial → API design)',
        'Adaptación rápida a nueva tecnología: aprendiste algo nuevo y lo aplicaste con impacto',
        'Mejoraste un proceso que ya funcionaba usando información nueva',
        'Ejecución: no solo aprendiste, sino que lo materializaste en resultado medible',
    ]
    for e in eval_3b:
        add_bullet(doc, e, size=9)
    add_para(doc, 'Frases que disparan señales positivas:', bold=True, size=10)
    signals_3 = [
        '"I helped because it was the right thing to do, not because it was my job"',
        '"I noticed he/she had no one to turn to — I made sure that changed"',
        '"I read [book/field outside software] and extracted a key insight that I applied to [work problem]"',
        '"The process was already working, but I saw that ‘working\' wasn\'t the same as \'best possible\'"',
        '"That learning changed how I think about [your domain] forever"',
    ]
    for s in signals_3:
        add_bullet(doc, s, size=9)
    add_para(doc, 'Red flags a EVITAR:', bold=True, size=10)
    red_3 = [
        'Historias de "caring" que en realidad son de tu logro personal, no del otro',
        'Hablar solo de tecnología que aprendiste dentro de tu stack (no demuestra curiosidad fuera del campo)',
        'Aplicar el aprendizaje sin resultado medible — "lo apliqué y fue interesante" no es suficiente',
        'Caring stories en "nosotros" sin clarificar qué fue TU aportación específica',
    ]
    for r in red_3:
        add_bullet(doc, r, size=9)

    add_section_title(doc, 'PREGUNTAS BEHAVIORAL PROBABLES', color_hex='BF8F00')
    questions_3 = [
        'ACTIVATE CARING:',
        '"Tell me about an experience you had in dedicating your time and expertise to help others."',
        '"Tell me about a time you had to balance your goal with a teammate\'s goal."',
        '"Tell me about a time you spent time getting to know a colleague better and why."',
        '"Tell me about a time you offered your voice in support of someone else\'s unpopular idea."',
        '',
        'STUDY & DRAW INSPIRATION:',
        '"Tell me about a time you used new information or insight to improve a product/process that was working well."',
        '"Tell me about the most interesting thing you have learned in the past month."',
        '"Tell me about a time you consumed information not relevant to your field that influenced your work."',
    ]
    for q in questions_3:
        if q == '':
            doc.add_paragraph()
        elif q.endswith(':'):
            add_para(doc, q, bold=True, size=10, color=RGBColor(0xBF, 0x8F, 0x00))
        else:
            add_bullet(doc, q)

    add_section_title(doc, 'STAR STORIES PREPARADAS', color_hex='BF8F00')

    add_star_story(
        doc,
        'STORY 1 — Mentorship Junior Developer (Thomson Reuters) [Activate Caring]',
        'Ayudaste sin obligación, tu impacto fue personal y profesional, no solo técnico.',
        {
            'S': 'En Thomson Reuters noté que un junior de otro equipo estaba batallando. Su equipo no lo apoyaba y estaba completamente bloqueado.',
            'T': 'No era mi obligación ayudar — no era mi proyecto, no era mi equipo.',
            'A': 'Me senté con él a explorar el problema. Debuggeamos juntos hasta encontrar la solución. Después empezó a venir conmigo cada vez que se bloqueaba — se convirtió en un mentorship informal donde le enseñé clean code, code reviews y buenas prácticas.',
            'R': 'Lo vi crecer y volverse más confiado. Para mí no fue trabajo extra — fue simplemente lo correcto.',
        },
        {
            'S': 'At Thomson Reuters I noticed a junior developer on another team was struggling. His team wasn\'t supporting him and he was completely blocked.',
            'T': 'I had no obligation to help — it wasn\'t my project or my team.',
            'A': 'I sat down with him and explored the problem together. We debugged it until we found the solution. After that he started coming to me whenever he was blocked — it became an informal mentorship where I shared clean code principles, code review techniques, and best practices.',
            'R': 'I saw him grow and become more confident over time. For me it wasn\'t extra work — it was just the right thing to do.',
        },
        color_hex='BF8F00',
    )

    add_star_story(
        doc,
        'STORY 2 — AI Code Review Agent (Thomson Reuters) [Study & Draw Inspiration]',
        'Aprendiste sobre AI/LLMs y lo aplicaste para mejorar un proceso que ya funcionaba.',
        {
            'S': 'En Thomson Reuters empecé a experimentar con Claude Code — un asistente AI de programación. Me fascinó cómo razona sobre patrones en grandes corpus de código.',
            'T': 'Nuestro proceso de code review ya funcionaba. Pero era inconsistente — juniors recibían feedback diferente dependiendo de quién los revisara.',
            'A': 'Apliqué lo que aprendí sobre AI pattern recognition: construi un agente custom que analizaba todos nuestros comentarios de PRs pasados, extraia patrones específicos del equipo, y automáticamente revisaba nuevos PRs antes de que cualquier humano los tocara.',
            'R': 'Reviews más rápidas y consistentes. Junior devs con feedback instantáneo y estandarizado. El equipo gastaba menos tiempo en comentarios repetitivos y más en discusiones arquitectónicas.',
        },
        {
            'S': 'At Thomson Reuters I started experimenting with Claude Code — an AI coding assistant. What fascinated me was how LLMs reason about patterns across large code corpora.',
            'T': 'Our code review process was working fine. But reviews were inconsistent — junior developers received different feedback depending on who reviewed them.',
            'A': 'I applied what I learned about AI pattern recognition: I built a custom agent that analyzed all our past PR comments, extracted team-specific patterns, and automatically reviewed new PRs before any human touched them.',
            'R': 'Reviews became faster and more consistent. Junior developers got standardized instant feedback. The team spent less time on repetitive comments and more on meaningful architectural discussions.',
        },
        color_hex='BF8F00',
    )

    add_star_story(
        doc,
        'STORY 3 — Design of Everyday Things → API Design [Study & Draw (outside field)]',
        'Leiste un libro fuera de software y lo convertiste en un cambio de diseño con resultado medible.',
        {
            'S': 'Leí "The Design of Everyday Things" de Don Norman — un libro de diseño industrial, no de software.',
            'T': 'Buscaba inspiración fuera de la tecnología para mejorar cómo diseñábamos APIs.',
            'A': 'El concepto de "affordances" — los objetos deben señalar cómo se usan — lo apliqué al diseño de APIs REST. En lugar de POST /getBookCover, usamos GET /books/{id}/cover. La API se vuelve self-documenting y reduce fricción de integración.',
            'R': 'Otros equipos adoptaron el patrón. Redujimos tickets de soporte relacionados a APIs en ~30%.',
        },
        {
            'S': 'I read "The Design of Everyday Things" by Don Norman — a book about industrial design, not software.',
            'T': 'I was looking for inspiration outside technology to improve how we designed our APIs for other teams.',
            'A': 'The concept of "affordances" — objects should signal how they\'re used — I applied directly to REST API design. Instead of POST /getBookCover, we used GET /books/{id}/cover. The API became self-documenting and reduced integration friction.',
            'R': 'Other teams adopted the pattern organically. We reduced support tickets related to API confusion by ~30%.',
        },
        color_hex='BF8F00',
    )

    # System Design problems
    add_section_title(doc, 'PROBLEMAS DE SYSTEM DESIGN PROBABLES', color_hex='BF8F00')
    add_para(doc, 'Audible se enfoca en content delivery, latencia y availability. Esperar problemas Audible-specific.', size=9)

    sd_problems = [
        ('Audible Social Site (Twitter-like feed)', 'Post Service + Kafka fan-out + Redis feed lists. Híbrido push/pull (celebridades pull, usuarios normales push). 10M DAU, eventual consistency.'),
        ('Design Audiobook Streaming (Audible core)', 'CDN (CloudFront) + S3 + DRM wrapper. Chunk-based delivery con pre-signed URLs. Progress sync via API cada X segundos.'),
        ('Design Bookmark Sync Across Devices', 'Event-based con timestamps. Last-write-wins o vector clocks. DynamoDB con userId como partition.'),
        ('Design a URL Shortener (bit.ly)', 'Base62 encoding de ID incremental. Redis cache con 80/20 split. 100:1 read/write ratio. Consistent hashing para sharding.'),
        ('Design a Notification Service', 'Kafka + Consumer workers por canal (email, push, SMS). Retry con exponential backoff. Dead letter queue.'),
        ('Design a Recommendation System', 'Collaborative filtering offline + online reranking. Redis cache de top N por usuario. A/B testing framework.'),
        ('Design a Chat System (Slack/Discord)', 'WebSockets para real-time. Cassandra para message history. Fan-out a channel members via Kafka.'),
        ('Design a Rate Limiter (distributed)', 'Redis + Lua scripts atómicos. Token bucket per user. Sliding window log vs fixed window trade-offs.'),
        ('Design a Search Autocomplete', 'Trie + prefix matching. Cache top N por prefix. Update asíncrono de weights.'),
        ('Design a File Sharing Service (Dropbox)', 'Chunking de files (4MB chunks). Metadata en SQL, blobs en S3. Delta sync para updates.'),
        ('Design YouTube/Netflix Video Streaming', 'Transcoding pipeline + adaptive bitrate (HLS/DASH). CDN para edge delivery. Metadata en distributed DB.'),
        ('Design Newsfeed (Facebook/Twitter)', 'Fan-out on write para 99% usuarios. Pull para influencers. Redis feeds con TTL.'),
    ]
    for name, desc in sd_problems:
        p = doc.add_paragraph()
        p.add_run('• ' + name + ': ').bold = True
        r = p.add_run(desc)
        r.font.size = Pt(9)

    add_section_title(doc, 'FRAMEWORK SYSTEM DESIGN (45-60 min)', color_hex='BF8F00')
    framework = [
        '1. REQUIREMENTS (5 min): Functional + Non-functional. Las 4 clarifying questions: DAU, consistency, features scope, constraints',
        '2. ESTIMATION (5 min): DAU → QPS → Storage → Bandwidth. Usa orden de magnitud, no exactitud',
        '3. API DESIGN (5 min): REST endpoints con request/response. POST /posts, GET /feed',
        '4. HIGH-LEVEL DESIGN (10 min): Diagrama de cajas. Client → LB → Services → DBs',
        '5. DATA MODEL (5 min): Schema principal. SQL vs NoSQL trade-off explícito',
        '6. DEEP DIVE (15 min): Chaitra guía — cache strategy, fan-out, sharding, fault tolerance',
        '7. TRADE-OFFS (5 min): Articula 2-3 decisiones con cases for/against',
        '8. SCALE / BOTTLENECKS (5 min): ¿Qué rompe primero? ¿Cómo lo detectas? Monitoring/alerts',
    ]
    for line in framework:
        add_bullet(doc, line, size=9)

    add_section_title(doc, 'CONCEPTOS CLAVE A MENCIONAR', color_hex='BF8F00')
    concepts = [
        'CAP theorem: escoge Availability + Partition tolerance (AP) para feed; Consistency + Partition (CP) para billing',
        'Consistent hashing: para sharding horizontal sin hotspots',
        'Bloom filter: para check de existencia rápido antes de hit a DB',
        'CDN (CloudFront, Akamai): edge caching para static content y audio',
        'Message queue (Kafka, SQS): desacoplar servicios, fan-out, at-least-once delivery',
        'Read replicas + write primary: leer de réplicas para scale de reads',
        'Cache strategies: cache-aside, write-through, write-back. TTL + invalidation',
        'Circuit breaker pattern: fault tolerance en llamadas downstream',
        'Rate limiting: token bucket en API gateway (Kong, AWS API Gateway)',
        'Monitoring: metrics (CloudWatch), distributed tracing (X-Ray), alerts',
    ]
    for c in concepts:
        add_bullet(doc, c, size=9)

    # EJEMPLOS DE CÓDIGO para Interview 3 (System Design con snippets)
    add_section_title(doc, 'EJEMPLOS DE CODIGO CLAVE (Java)', color_hex='BF8F00')
    add_para(doc, 'Aunque es System Design, Chaitra puede pedirte código de ciertos componentes. Estos son snippets cortos que demuestran patrones distribuidos en Java.', size=9)

    # API contract example
    api_code = '''// REST endpoints para Audible Social Site

@RestController
@RequestMapping("/api/v1")
public class PostController {

    @Autowired
    private PostService postService;

    @Autowired
    private FeedService feedService;

    @PostMapping("/posts")
    public ResponseEntity<PostResponse> createPost(
            @RequestBody @Valid CreatePostRequest request,
            @AuthenticationPrincipal User user) {
        Post post = postService.create(user.getId(),
                                       request.getContent(),
                                       request.getBookReference());
        return ResponseEntity.status(HttpStatus.CREATED)
                .body(PostResponse.from(post));
    }

    @GetMapping("/feed")
    public ResponseEntity<FeedResponse> getFeed(
            @AuthenticationPrincipal User user,
            @RequestParam(defaultValue = "20") int limit,
            @RequestParam(required = false) String cursor) {
        FeedPage feed = feedService.getFeed(user.getId(), limit, cursor);
        return ResponseEntity.ok(FeedResponse.from(feed));
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run('→ REST API Contract (Audible Social Site)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
    add_code(doc, api_code, size=8)

    # Fan-out service with Kafka
    fanout_code = '''// Fan-out Service — escribe post_id en el feed de cada follower via Redis

@Service
public class FanOutService {

    @Autowired
    private FollowerRepository followerRepo;

    @Autowired
    private RedisTemplate<String, String> redis;

    private static final int MAX_FEED_SIZE = 1000;
    private static final long CELEBRITY_THRESHOLD = 10_000L;

    @KafkaListener(topics = "new-posts")
    public void handleNewPost(NewPostEvent event) {
        long followerCount = followerRepo.countByUserId(event.getAuthorId());

        // Híbrido: push para users normales, pull para celebridades
        if (followerCount >= CELEBRITY_THRESHOLD) {
            // Pull model: solo guardar en user timeline, feed service hace join on read
            return;
        }

        List<String> followers = followerRepo.findFollowerIds(event.getAuthorId());
        for (String followerId : followers) {
            String feedKey = "feed:" + followerId;
            redis.opsForList().leftPush(feedKey, event.getPostId());
            redis.opsForList().trim(feedKey, 0, MAX_FEED_SIZE - 1);
            redis.expire(feedKey, Duration.ofDays(7));
        }
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run('→ Fan-out Service (Kafka consumer + Redis feed)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
    add_code(doc, fanout_code, size=8)

    # Circuit breaker pattern
    circuit_code = '''// Circuit Breaker — fault tolerance al llamar servicios downstream

@Component
public class RecommendationClient {

    private final CircuitBreaker breaker;
    private final RestTemplate http;

    public RecommendationClient(CircuitBreakerRegistry registry, RestTemplate http) {
        this.breaker = registry.circuitBreaker("recommendations");
        this.http = http;
    }

    public List<String> getRecommendations(String userId) {
        return breaker.executeSupplier(() ->
            http.getForObject(
                "http://recommendations-svc/users/" + userId + "/recs",
                RecommendationResponse.class
            ).getItems()
        );
    }

    // Fallback cuando el circuito está abierto
    public List<String> getRecommendationsWithFallback(String userId) {
        try {
            return getRecommendations(userId);
        } catch (CallNotPermittedException e) {
            return getDefaultRecommendations();
        }
    }

    private List<String> getDefaultRecommendations() {
        return List.of("popular_1", "popular_2", "popular_3");
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run('→ Circuit Breaker pattern (resiliencia)')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
    add_code(doc, circuit_code, size=8)

    add_section_title(doc, 'PREGUNTAS PARA CHAITRA', color_hex='BF8F00')
    questions_for_them_3 = [
        '"What system design challenges has your team faced recently that were the most interesting?"',
        '"How does Audible handle scale during peak listening times?"',
        '"What is the team culture around knowledge sharing and mentorship?"',
        '"What technologies excites you the most that you work with at Audible?"',
    ]
    for q in questions_for_them_3:
        add_bullet(doc, q)

    add_page_break(doc)

    # ============================================================
    # INTERVIEW 4 - POORVA KARUNAKARAN
    # ============================================================
    add_heading(doc, 'INTERVIEW 4: Poorva Karunakaran (SDE II)', level=1, color=RGBColor(0xC0, 0x00, 0x00))
    add_para(doc, '1:00-2:00pm  |  Coding (Logical & Maintainable) + Imagine & Invent Before They Ask', bold=True, size=11)
    doc.add_paragraph()

    add_section_title(doc, 'PRINCIPIO: Imagine & Invent Before They Ask', color_hex='C00000')
    add_para(doc, '"We embrace ambiguity and tolerate and demystify complexity. We make customers feel the power before they have to think to ask for it."', size=9)
    add_para(doc, 'Lo que evalúan:', bold=True, size=10)
    eval_4 = [
        'Proactividad sin instrucción: viste un problema y actuaste antes de que alguien te lo pidiera',
        'Abraza la ambigüedad: tomas acción aunque no tengas todos los datos, defines el problema tú mismo',
        'Simplificación: redujiste complejidad para el usuario o el equipo (no solo optimización técnica)',
        'Mejoraste algo que ya funcionaba porque viste cómo podía ser mejor',
        'Aprendiste de un fracaso y lo documentaste como éxito para otros',
    ]
    for e in eval_4:
        add_bullet(doc, e, size=9)
    add_para(doc, 'Frases que disparan señales positivas:', bold=True, size=10)
    signals_4 = [
        '"Nobody asked me to fix this — I just saw the opportunity and moved"',
        '"I didn\'t wait for perfect information. I defined the problem, ran a quick experiment, and iterated"',
        '"I made it invisible to the user — they didn\'t even notice the change, they just experienced it better"',
        '"I documented what I learned so three other teams didn\'t have to repeat the same mistakes"',
        '"The process was working fine, but \'fine\' wasn\'t good enough for me"',
    ]
    for s in signals_4:
        add_bullet(doc, s, size=9)
    add_para(doc, 'Red flags a EVITAR:', bold=True, size=10)
    red_4 = [
        'Historias donde esperaste permiso o instrucciones antes de actuar',
        'Mejoras que no tienen impacto medible en el usuario o el equipo',
        'Inventar algo que ya existía — no mostraste que investigaste alternativas primero',
        'Ambigüedad que te paralizó en lugar de invitarte a explorar',
    ]
    for r in red_4:
        add_bullet(doc, r, size=9)

    add_section_title(doc, 'PREGUNTAS BEHAVIORAL PROBABLES', color_hex='C00000')
    questions_4 = [
        '"Tell me about a time you took something that was working just fine and still found a way to improve it."',
        '"Tell me about a time you were able to see around the corner to meet a customer need before they requested it."',
        '"Tell me about a time you delivered an outcome that failed but you deemed it a success due to learnings."',
        '"Tell me about a time you simplified a complex customer experience."',
        '"Tell me about a time you embraced ambiguity to deliver results."',
    ]
    for q in questions_4:
        add_bullet(doc, q)

    add_section_title(doc, 'STAR STORIES PREPARADAS', color_hex='C00000')

    add_star_story(
        doc,
        'STORY 1 \u2014 AI Code Review Agent (Thomson Reuters) [Improved What Was Working]',
        'Nadie te pidió que lo hicieras. Viste el problema, definiste la solución, y ejecutaste con datos.',
        {
            'S': 'En Thomson Reuters empezamos un proyecto nuevo y hubo fricción inmediata — code reviews tomaban demasiado porque no había criterios claros de aceptación.',
            'T': 'Nadie me pidió arreglar esto. Yo vi el problema y decidí actuar.',
            'A': 'Corrí polls con el equipo para entender sus dolores. Usé Claude Code para analizar todos los comentarios pasados de PRs y extraer patrones específicos del equipo. Con esos datos redacté criterios de aceptación. Y fui más allá: construí un agente custom que automáticamente revisaba nuevos PRs antes de cualquier revisor humano.',
            'R': 'Reviews más rápidas y consistentes. Feedback instantáneo para junior devs. El equipo pasó de comentarios repetitivos a discusiones arquitectónicas de mayor valor.',
        },
        {
            'S': 'At Thomson Reuters we started a new project and immediately had friction — code reviews were taking too long because there were no clear acceptance criteria.',
            'T': 'Nobody asked me to fix this. I just saw the problem and decided to act.',
            'A': 'I ran polls with the team to understand their pain points. Then I used Claude Code to analyze all our past PR comments and extract team-specific patterns. From that data I drafted our acceptance criteria. But I pushed further — I built a custom AI code review agent that automatically checked new PRs against those rules before any human reviewer touched the code.',
            'R': 'Code reviews became faster and more consistent. Junior developers got instant feedback. The team moved from repetitive comments to meaningful architectural discussions.',
        },
        color_hex='C00000',
    )

    add_star_story(
        doc,
        'STORY 2 \u2014 Migración GCP Transparente (Kubo) [Invented, Simplified, Invisible to User]',
        'Encontraste una solución no-obvia que simplificó radicalmente el trabajo sin que el usuario lo notara.',
        {
            'S': 'El monolito en Kubo Financiero se caía frecuentemente por falta de espacio en disco. La solución obvia era migrar a la nube — pero requería semanas de cambios en el frontend JSP.',
            'T': 'Mi tarea era resolver los crashes con el equipo sin destrozar la aplicación.',
            'A': 'Imaginé una solución invisible: migrar los archivos a GCP pero mapear los paths de la nube a los mismos paths locales que el sistema ya esperaba. Los archivos se descargaban on-demand de forma transparente — cero cambios al código del frontend.',
            'R': 'Crashes resueltos con un mínimo de cambios. Uso de recursos del servidor reducido significativamente. El frontend ni se enteró de la migración.',
        },
        {
            'S': 'The monolith at Kubo Financiero was crashing frequently due to disk space issues. The obvious solution was to migrate to the cloud — but our JSP-based frontend would have required weeks of changes.',
            'T': 'My task was to solve the crashes with my team without disrupting the live application.',
            'A': 'I imagined an invisible solution: migrate files to GCP but map cloud paths to the same local paths the system already expected. Files would download on-demand transparently — zero changes to frontend code.',
            'R': 'Crashes resolved with minimal code changes. Server resource usage reduced significantly. The frontend never knew the migration happened.',
        },
        color_hex='C00000',
    )

    # Problemas Logical & Maintainable
    add_section_title(doc, 'PROBLEMAS DE CODING PROBABLES (Logical & Maintainable)', color_hex='C00000')
    add_para(doc, 'Poorva busca código limpio, extensible, bien nombrado, con buenos abstractions. Menos foco en optimización extrema, más en mantenibilidad.', size=9)

    lm_problems = [
        ('GetCommunity (Tree filter)', 'DFS post-order recursivo. Si nodo matches → crea ResultNode con matching children. Si no → propaga children hacia arriba. Extensible: usar Predicate<T> en lugar de string param.'),
        ('File Search with Rules', 'Composite pattern: FileRule interface, FileRuleAnd/Or toman varargs. Strategy pattern para matchers (extension, size, name). Recursión en directorios.'),
        ('Design a Deck of Cards (Blackjack)', 'Enum Suit y Rank. Card class. Deck con shuffle/draw. Player abstracto. BlackjackHand extends Hand. Factory para diferentes juegos.'),
        ('Design a Chess Game', 'Piece abstract class. Cada pieza implementa canMove(). Board 8x8. Game tracks turns, checkmate detection. Observer pattern para UI updates.'),
        ('Design a Logger / Observer Pattern', 'Log levels enum. Appender interface (Console, File, Remote). Chain of responsibility para filtering.'),
        ('Design LRU Cache (clean impl)', 'HashMap<K, Node>. DoublyLinkedList con head/tail sentinels. moveToFront() helper. O(1) get/put.'),
        ('Snake Game', 'Board 2D. Snake = Deque<Position>. Food random spawn. Collision detection: self o wall. Direction enum.'),
        ('Design a Meeting Room Scheduler', 'Interval merging. Priority queue por end time. Min rooms needed = max overlapping intervals.'),
        ('Design a File System (OOP)', 'Abstract FileSystemEntry. File y Directory extend. Composite pattern para tree structure. Visitor pattern para traversal.'),
        ('Design a Tic-Tac-Toe (OOP)', 'Board 3x3. Player con marker. Track rowSum, colSum, diag. checkWin() en O(1) después de cada move.'),
        ('Implement a Trie', 'TrieNode con Map<Char, TrieNode>. insert, search, startsWith. Útil para autocomplete, spell check.'),
        ('Design a Phone Directory', 'Interface + HashMap impl. get, release, check. Alternative: queue de available numbers.'),
    ]
    for name, desc in lm_problems:
        p = doc.add_paragraph()
        p.add_run('• ' + name + ': ').bold = True
        r = p.add_run(desc)
        r.font.size = Pt(9)

    add_section_title(doc, 'PRINCIPIOS SOLID A ARTICULAR EXPLICITAMENTE', color_hex='C00000')
    solid = [
        'S - Single Responsibility: cada clase tiene UNA razón para cambiar',
        'O - Open/Closed: abierto para extensión, cerrado para modificación (usar interfaces)',
        'L - Liskov Substitution: subclases deben poder reemplazar a la base sin romper',
        'I - Interface Segregation: interfaces pequeñas y específicas, no "god interfaces"',
        'D - Dependency Inversion: depende de abstracciones, no de implementaciones concretas',
    ]
    for s in solid:
        add_bullet(doc, s, size=9)

    add_section_title(doc, 'DESIGN PATTERNS CLAVE A MENCIONAR', color_hex='C00000')
    patterns = [
        'Strategy: algoritmos intercambiables (parking assignment, file matching)',
        'State: comportamiento cambia según estado (elevator, vending machine)',
        'Observer: notificaciones asíncronas (library returns, stock alerts)',
        'Factory: creación de objetos compleja (deck of cards variants)',
        'Composite: estructuras tipo árbol (filesystem, HTML DOM)',
        'Decorator: añadir responsabilidades dinámicamente (UI, streams)',
        'Singleton: una sola instancia (logger, config, DB connection pool)',
        'Chain of Responsibility: pasar request a través de handlers (middleware, log filters)',
        'Builder: construcción compleja paso a paso (StringBuilder, SQL query)',
        'Adapter: interfaz incompatible (legacy code integration)',
    ]
    for pat in patterns:
        add_bullet(doc, pat, size=9)

    add_section_title(doc, 'PROTOCOLO PARA CODE LOGICAL & MAINTAINABLE', color_hex='C00000')
    protocol_4 = [
        '1. CLARIFY: scope, requirements, constraints',
        '2. CLASS DIAGRAM primero: entidades y relaciones antes de código',
        '3. NAMING: clases como sustantivos (ParkingLot, not parking_handler), métodos como verbos (assignSpot)',
        '4. INTERFACES antes de implementaciones: programa a la abstracción',
        '5. MENCIONA SOLID mientras codeas: "this follows Open/Closed because..."',
        '6. EDGE CASES: null inputs, empty collections, concurrent access',
        '7. EXTENSIBILIDAD: "if the requirements change to X, we would just add..."',
    ]
    for line in protocol_4:
        add_bullet(doc, line, size=9)

    add_section_title(doc, 'CODIGO EJEMPLO (Java)', color_hex='C00000')
    add_para(doc, 'Ejemplos de código limpio y mantenible. El clean code en Java es lo que Poorva busca específicamente.', size=9)

    add_code_example(
        doc,
        'Balanced Binary Tree (recursión limpia con -1 sentinel)',
        '/home/pinky/proyects/leetcode/src/BalanceTreeSolution.java',
        'Patrón: DFS post-order. Usa -1 como sentinel para unbalanced, evita O(n²). Técnica clave para el GetCommunity problem.'
    )

    add_code_example(
        doc,
        'Merge Two Sorted Lists (dummy node pattern)',
        '/home/pinky/proyects/leetcode/src/MergeTwoSortedListsSolution.java',
        'Patrón: dummy node sentinel para evitar casos especiales en linked lists. Código extremadamente limpio.'
    )

    # GetCommunity - the actual problem from Poorva's example
    getcommunity_code = '''// ============================================
// GetCommunity — Tree filter con Predicate
// Patrón: DFS post-order + Composite + Strategy
// ============================================
public class Employee {
    String id;
    String name;
    Map<String, String> properties;
    List<Employee> directReports;
}

public class ResultNode {
    String id;
    String name;
    List<ResultNode> matchingChildren = new ArrayList<>();

    public ResultNode(String id, String name) {
        this.id = id;
        this.name = name;
    }
}

public class CommunityFinder {

    /**
     * Extensible: acepta cualquier Predicate<Employee>.
     * Open/Closed: nuevas condiciones sin tocar esta clase.
     */
    public List<ResultNode> getCommunity(Employee root, Predicate<Employee> matcher) {
        if (root == null) {
            return Collections.emptyList();
        }
        return dfs(root, matcher);
    }

    /**
     * Post-order DFS: procesa hijos antes que el padre.
     * - Si el nodo hace match: retorna [ResultNode con matchingChildren]
     * - Si NO hace match: se "vuelve transparente" y propaga los matches de hijos
     */
    private List<ResultNode> dfs(Employee node, Predicate<Employee> matcher) {
        List<ResultNode> childMatches = new ArrayList<>();

        if (node.directReports != null) {
            for (Employee child : node.directReports) {
                childMatches.addAll(dfs(child, matcher));
            }
        }

        if (matcher.test(node)) {
            ResultNode result = new ResultNode(node.id, node.name);
            result.matchingChildren = childMatches;
            return List.of(result);
        }

        // node no hace match: propaga hijos hacia arriba
        return childMatches;
    }
}

// ============================================
// Uso: múltiples condiciones sin cambiar la API
// ============================================
CommunityFinder finder = new CommunityFinder();

// Caso 1: filtrar por un property value
Predicate<Employee> byCity = e -> "CDMX".equals(e.properties.get("city"));
List<ResultNode> cdmxCommunity = finder.getCommunity(root, byCity);

// Caso 2: filtro compuesto (AND)
Predicate<Employee> seniorInCdmx = byCity
    .and(e -> "SENIOR".equals(e.properties.get("level")));
List<ResultNode> seniorCdmx = finder.getCommunity(root, seniorInCdmx);

// Caso 3: OR
Predicate<Employee> javaOrKotlin = e ->
    "java".equals(e.properties.get("stack")) ||
    "kotlin".equals(e.properties.get("stack"));
'''
    add_code_example(
        doc,
        'GetCommunity — el problema específico de Poorva',
        '',
        'Solución completa con Predicate<Employee> para extensibilidad. Aplica SOLID (Open/Closed) y Strategy pattern.'
    )
    add_code(doc, getcommunity_code, size=8)

    # Clean Observer pattern example
    observer_code = '''// ============================================
// Observer Pattern — típico OOD question
// Maintainable: subjects y observers desacoplados
// ============================================
public interface Observer<T> {
    void onEvent(T event);
}

public interface Subject<T> {
    void subscribe(Observer<T> observer);
    void unsubscribe(Observer<T> observer);
    void publish(T event);
}

public class EventBus<T> implements Subject<T> {
    private final List<Observer<T>> observers = new CopyOnWriteArrayList<>();

    @Override
    public void subscribe(Observer<T> observer) {
        observers.add(observer);
    }

    @Override
    public void unsubscribe(Observer<T> observer) {
        observers.remove(observer);
    }

    @Override
    public void publish(T event) {
        for (Observer<T> observer : observers) {
            try {
                observer.onEvent(event);
            } catch (Exception e) {
                // one bad observer shouldn\\'t break others
                log.warn("Observer failed", e);
            }
        }
    }
}

// Uso:
EventBus<OrderPlaced> bus = new EventBus<>();
bus.subscribe(event -> emailService.sendConfirmation(event));
bus.subscribe(event -> inventoryService.reserve(event));
bus.subscribe(event -> analyticsService.track(event));
bus.publish(new OrderPlaced("order-123", userId));
'''
    add_code_example(
        doc,
        'Observer Pattern (thread-safe, typical OOD question)',
        '',
        'CopyOnWriteArrayList para thread-safety. Try/catch por observer para aislamiento de fallos. Patrón base para Library Mgmt System, Stock Alerts.'
    )
    add_code(doc, observer_code, size=8)

    # Parking Lot Strategy example
    parking_code = '''// ============================================
// Parking Lot — típico OOD (Strategy pattern)
// ============================================
public enum VehicleSize { MOTORCYCLE, COMPACT, LARGE }

public abstract class Vehicle {
    protected String plate;
    protected VehicleSize size;
    public abstract int getSpotsNeeded();
}

public class Car extends Vehicle {
    public Car(String plate) {
        this.plate = plate;
        this.size = VehicleSize.COMPACT;
    }
    @Override public int getSpotsNeeded() { return 1; }
}

public class Bus extends Vehicle {
    public Bus(String plate) {
        this.plate = plate;
        this.size = VehicleSize.LARGE;
    }
    @Override public int getSpotsNeeded() { return 5; }
}

// Strategy: diferentes políticas de asignación
public interface SpotAssignmentStrategy {
    Optional<List<ParkingSpot>> findSpots(Vehicle vehicle, List<ParkingSpot> available);
}

public class NearestSpotStrategy implements SpotAssignmentStrategy {
    @Override
    public Optional<List<ParkingSpot>> findSpots(Vehicle v, List<ParkingSpot> available) {
        // lógica: spots contiguos más cercanos a la entrada
        return findContiguousSpots(available, v.getSpotsNeeded());
    }
    private Optional<List<ParkingSpot>> findContiguousSpots(List<ParkingSpot> s, int n) {
        for (int i = 0; i <= s.size() - n; i++) {
            if (allAvailable(s, i, n)) {
                return Optional.of(s.subList(i, i + n));
            }
        }
        return Optional.empty();
    }
    private boolean allAvailable(List<ParkingSpot> s, int start, int n) {
        for (int i = start; i < start + n; i++) {
            if (!s.get(i).isAvailable()) return false;
        }
        return true;
    }
}

public class ParkingLot {
    private final List<ParkingSpot> spots;
    private final SpotAssignmentStrategy strategy;  // injected

    public ParkingLot(List<ParkingSpot> spots, SpotAssignmentStrategy strategy) {
        this.spots = spots;
        this.strategy = strategy;
    }

    public Optional<Ticket> park(Vehicle vehicle) {
        List<ParkingSpot> available = spots.stream()
            .filter(ParkingSpot::isAvailable)
            .collect(Collectors.toList());
        return strategy.findSpots(vehicle, available)
            .map(assigned -> {
                assigned.forEach(s -> s.occupy(vehicle));
                return new Ticket(vehicle, assigned, Instant.now());
            });
    }
}
'''
    add_code_example(
        doc,
        'Parking Lot (Strategy pattern, extensible)',
        '',
        'OOD clásico. Strategy para políticas de asignación intercambiables. Dependency Injection. Open/Closed.'
    )
    add_code(doc, parking_code, size=8)

    add_section_title(doc, 'PREGUNTAS PARA POORVA', color_hex='C00000')
    questions_for_them_4 = [
        '"How does your team approach code reviews and maintaining code quality?"',
        '"What is the most interesting problem youre currently solving at Audible?"',
        '"How do you balance technical debt with delivering new features?"',
        '"What do you think differentiates a great SDE II from a good one?"',
    ]
    for q in questions_for_them_4:
        add_bullet(doc, q)

    add_page_break(doc)

    # ============================================================
    # TIPS GENERALES + CHECKLIST
    # ============================================================
    add_heading(doc, 'TIPS GENERALES PARA TODO EL DIA', level=1, color=RGBColor(0x40, 0x40, 0x40))

    add_section_title(doc, 'EL DIA ANTES')
    pre_day = [
        'Duerme bien — mínimo 7 horas. El cerebro es el instrumento principal.',
        'Prepara el setup: laptop cargada, cable de backup, internet cableado si es posible',
        'Zoom pre-instalado. Test del micrófono, cámara y audio.',
        'Ten el cheatsheet abierto pero NO lo leas durante la entrevista — solo para glance rápidos',
        'Prepara agua, café/té, snacks cerca',
        'Baño antes de cada sesión, tienes 5 min entre entrevistas',
    ]
    for t in pre_day:
        add_bullet(doc, t)

    add_section_title(doc, 'DURANTE LA ENTREVISTA')
    during = [
        'Saluda cálidamente, sonríe. Primera impresión cuenta.',
        'Escucha la pregunta COMPLETA. No empieces a responder antes de que termine.',
        'Si no entiendes algo, pide clarificación — nunca asumas.',
        'Piensa en voz alta. El silencio prolongado es enemigo.',
        'Usa "Let me think about this for a moment" si necesitas pausar (2-3 seg OK).',
        'En STAR: máximo 2 minutos por historia. No te extiendas.',
        'Usa "I" (yo), no "we" (nosotros) — ellos evalúan TU contribución.',
        'Si cometes un error en código, NO entres en pánico: "I notice I have a bug here, let me fix it".',
        'Al final: "Do you have any questions for me?" — SIEMPRE ten 2-3 preguntas listas.',
    ]
    for t in during:
        add_bullet(doc, t)

    add_section_title(doc, 'SEÑALES QUE QUIEREN VER')
    signals = [
        '"Let me clarify first..." → demuestras pensamiento estructurado',
        '"The trade-off here is..." → senior-level thinking',
        '"What if I sort the array? That gives us O(n log n)..." → exploración de alternativas',
        '"Edge case: what if the input is empty?" → rigor',
        '"I\'d want to add monitoring on X because..." → mentalidad operacional',
        '"I owned this decision because..." → leadership',
        '"Looking back, I would have done Y differently" → self-awareness',
    ]
    for s in signals:
        add_bullet(doc, s, size=9)

    add_section_title(doc, 'RED FLAGS A EVITAR')
    red_flags = [
        'Empezar a codificar sin clarificar',
        'Quedarse en silencio más de 10 segundos',
        'Decir "this is easy" o "obvious" — suena arrogante',
        'Pelear con el entrevistador en trade-offs — escucha y adapta',
        'Historias STAR en "we" todo el tiempo — no queda claro tu contribución',
        'No tener preguntas al final',
        'Mencionar salario o beneficios en técnicas (solo con el recruiter)',
        'Criticar empleadores anteriores',
    ]
    for rf in red_flags:
        add_bullet(doc, rf, size=9)

    add_section_title(doc, 'FRASES EN INGLES LISTAS PARA USAR')
    phrases = [
        'Clarificar: "Just to make sure I understand correctly, you\'re asking..."',
        'Pensar: "Let me think through this out loud..."',
        'Trade-off: "There are a few approaches here. The first would be... The trade-off is..."',
        'Error: "I notice there\'s a bug in my code, let me trace through it..."',
        'Cerrar: "So to summarize, my approach is X because Y, with O(n) time and O(n) space"',
        'Preguntas: "What does success look like in this role?" / "What are the biggest challenges the team is facing?"',
    ]
    for p in phrases:
        add_bullet(doc, p, size=9)

    add_section_title(doc, 'DATOS PERSONALES RAPIDOS DE TI')
    bio = [
        'Nombre: Alejandro Hernandez Loza',
        'Titulo: SR Software Engineer — 10+ años total, 5+ en senior roles',
        'Stack: Java, Spring Boot, Full Stack, JavaScript, Cloud (AWS/GCP), Docker, K8s, Microservices',
        'Experiencia reciente: Thomson Reuters (senior dev + AI experiments)',
        'Antes: Globant (migrations), Kubo Financiero (monolith rescue), Clickonero (Hot Sale incidents), Presidencia (CMS, terremoto)',
        'Ubicación: CDMX (remoto preferido)',
    ]
    for b in bio:
        add_bullet(doc, b)

    add_section_title(doc, 'PRONUNCIACION CORRECTA')
    pronunciation = [
        'Keycloak (no keycloack)',
        'Prometheus (no prometeus)',
        'Kubernetes (ku-ber-NET-is)',
        'PostgreSQL (POST-gres-QL)',
        'MySQL (my-S-Q-L o my-sequel)',
        'Nginx (engine-X)',
        'SQLite (S-Q-lite)',
    ]
    for p in pronunciation:
        add_bullet(doc, p, size=9)

    return doc


if __name__ == '__main__':
    doc = build_document()
    output_path = '/home/pinky/proyects/jobSearcher/docs/AUDIBLE_CHEATSHEET.docx'
    doc.save(output_path)
    print(f'Cheatsheet generado: {output_path}')
